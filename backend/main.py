from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from google import genai
from docx import Document
from pypdf import PdfReader
import os
import io
import json

load_dotenv()

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def root():
    return "hello"


@app.post("/upload_resume")
async def upload_resume(file: UploadFile = File(...)):
    contents = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(contents))
        resume_text = "\n".join([para.text for para in doc.paragraphs])

    elif filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(contents))
        resume_text = "\n".join(
            [page.extract_text() or "" for page in reader.pages]
        )

    else:
        return {"error": "Only .docx and .pdf files are supported"}

    prompt = f"""
    Analyze this resume.
    
    Return ONLY valid JSON.
    
    Format:
    {{
    "score": number,
    "strengths": [],
    "missing_skills": [],
    "suggestions": []
    }}
    
    Resume:
    {resume_text}
    """

    try:
        response = client.models.generate_content(
            model="gemini-3.1-flash-lite",
            contents=prompt
        )
        cleaned = (
            response.text
            .replace("```json", "")
            .replace("```", "")
            .strip()
        )

        data = json.loads(cleaned)

        return {"result": data}

    except Exception as e:
        return {"error": str(e)}


@app.post("/summarize")
async def summarize(file: UploadFile = File(...)):
    contents = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(contents))
        resume_text = "\n".join([para.text for para in doc.paragraphs])

    elif filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(contents))
        resume_text = "\n".join(
            [page.extract_text() or "" for page in reader.pages]
        )

    else:
        return {"error": "Only .docx and .pdf files are supported"}

    
    prompt = f"""
    Summarize this resume.
    
    Return ONLY valid JSON.
    
    Format:
    {{
    "title": "",
    "summary": ""
    }}
    
    Resume:
    {resume_text}
    """

    try:
        response = client.models.generate_content(
            model="gemini-3.1-flash-lite",
            contents=prompt,
        )
        cleaned = (
            response.text
            .replace("```json", "")
            .replace("```", "")
            .strip()
        )

        data = json.loads(cleaned)

        return {"result": data}

    except Exception as e:
        return {"error": str(e)}
